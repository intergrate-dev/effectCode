import json
import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR

sour_dir = './db-data/'
targ_dir = './db-data/out/'

def load_data(file):
    list_data = []
    NO = 1
    for l in open(file, 'r', encoding='utf-8'):
        if l.strip().startswith('CREATE TABLE') and l.index("(") > 0:
            data['t_name'] = extr_midd(l, '`')
        elif "PRIMARY KEY" in l:
            # data['pk_f'] = extr_midd(l, '`')
            pk_f.append(extr_midd(l, '`'))
        elif "ENGINE=" in l:
            data['t_desc'] = ''
            if 'COMMENT=' in l:
                data['t_desc'] = extr_midd(l, "'")
            create_table(data, list_data) 
            db_list.append([data['t_name'], data['t_desc']])
            list_data = []
            NO = 1
        elif l.strip().startswith("`") and "," in l:
            arr = proc_line(l)        
            list_data.append([str(NO), arr[0], arr[1], arr[2], arr[3], arr[4]]) 
            NO = NO + 1
           
def proc_line(line):
    arr[0] = extr_midd(line, "`")
    arr[1] = " "
    if "COMMENT" in line:
        arr[1] = extr_midd(line.split("COMMENT")[1], "'")
    arr[2] = line.strip().split(" ")[1]
    arr[3] = "否"
    if "NOT NULL" in line:
        arr[3] = "是"
    arr[4] = " "
    return arr

def extr_midd(str, flag):
    return str.split(flag)[1].split(flag)[0]

def write_table(file, head, content):
      file.write(head)
      file.write(content)

def create_table(data, records):
    # document.add_paragraph('表名：' + data['t_name'] + '\t描述：' + data['t_desc'])
    document.add_heading('表名：' + data['t_name'] + '\t描述：' + data['t_desc'], level=3)
    rows = 1
    cols = 6
    col_width_dic = {0: 1, 1: 3, 2: 4, 3: 2, 4: 2, 5: 2}
    title = [u'序号', u"字段", u"名称", u"数据类型", u"是否必填", u"注释"]
    gen_table(rows, cols, title, col_width_dic, records, set_table_cords)
    document.add_paragraph('\r')

def set_table_title(table, cols, col_width_dic, title):
    for col_num in range(cols):
        table.cell(0, col_num).width = Inches(col_width_dic[col_num])
        table.cell(0, col_num).text = title[col_num]
        table.cell(0, col_num).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_table_cords(table, records):
    for No, f, f_name, f_type, is_must, comment in records:
        # if f == data['pk_f']:
        if f in pk_f:
            f_name = '主键' + f
        row_cells = table.add_row().cells
        row_cells[0].text = str(No)
        row_cells[1].text = f
        row_cells[2].text = f_name
        row_cells[3].text = f_type
        row_cells[4].text = is_must
        row_cells[5].text = comment

def set_db_list_cords(table, records):
    i = 0
    for t_name, t_desc in records:
        i = i + 1
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = t_name
        row_cells[2].text = t_desc

def gen_table(rows, cols, title, col_width_dic, records, callback):
    table = document.add_table(rows, cols, style = "Table Grid")
    set_table_title(table, cols, col_width_dic, title)
    callback(table, records)

def gen_db_list():
    # document.add_paragraph('数据库表清单')
    document.add_heading('数据库表清单\t\t表总数：' + str(len(db_list)), level=2)
    rows = 1
    cols = 3
    col_width_dic = {0: 1, 1: 3, 2: 4}
    title = [u'序号', u'表名', u"描述"]      
    gen_table(rows, cols, title, col_width_dic, db_list, set_db_list_cords)
    document.add_paragraph('\r')  

if __name__ == '__main__':
    for f_name in os.listdir(sour_dir):
        data = {}
        pk_f = []
        arr = [0] * 5
        db_list = []
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        # if os.path.isfile(f_name):
        if '.sql' in f_name :
            document.add_heading('数据库表', level=2)
            load_data(sour_dir + f_name)
            if len(db_list) > 0:
                gen_db_list()
                document.save(targ_dir + f_name.split('.')[0] + '.docx')
                print(f_name + ", sql2word complete")

